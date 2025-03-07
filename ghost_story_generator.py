import streamlit as st
import requests
import json
import os
import sqlite3
from datetime import datetime
import uuid
import re
from docx import Document
from fpdf import FPDF
import tempfile
import time
import random

# Cấu hình API
API_KEYS = [
    "AIzaSyA6-W9fSgwDFSjf2i-gnirXwfaiah6M2zg",
    "AIzaSyC7lFa-0ZHvh09Hm4TtYfVc894UQXggLX0",
    "AIzaSyBWS4VBDtvkkbkvZyaawUbWAle4sXPS7YU",
    "AIzaSyAiAXsuJ9o1bCjaaRh2aUUYaTiZIFvU0Co",
    "AIzaSyDqESoT7B7CIkxfLBdC3DzbgjxbSVjq36o"
    # Thêm các API key khác vào đây
]
API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# Khởi tạo SQLite database
def init_db():
    conn = sqlite3.connect('ghost_stories.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS stories
                 (story_id TEXT PRIMARY KEY, outline TEXT, original_prompt TEXT, 
                  created_at TIMESTAMP, last_updated TIMESTAMP)''')
    c.execute('''CREATE TABLE IF NOT EXISTS chapters
                 (chapter_id TEXT PRIMARY KEY, story_id TEXT, chapter_number INTEGER,
                  content TEXT, created_at TIMESTAMP,
                  FOREIGN KEY (story_id) REFERENCES stories(story_id))''')
    c.execute('''CREATE TABLE IF NOT EXISTS chapter_versions
                 (version_id TEXT PRIMARY KEY, chapter_id TEXT, content TEXT,
                  created_at TIMESTAMP,
                  FOREIGN KEY (chapter_id) REFERENCES chapters(chapter_id))''')
    c.execute('''CREATE TABLE IF NOT EXISTS story_arcs
                 (arc_id TEXT PRIMARY KEY, story_id TEXT, arc_number INTEGER,
                  outline TEXT, created_at TIMESTAMP,
                  FOREIGN KEY (story_id) REFERENCES stories(story_id))''')
    conn.commit()
    conn.close()

# Gọi hàm khởi tạo database khi khởi động ứng dụng
init_db()

def get_db():
    return sqlite3.connect('ghost_stories.db')

# Biến đếm để theo dõi API key hiện tại
current_api_key_index = 0

def get_next_api_key():
    global current_api_key_index
    api_key = API_KEYS[current_api_key_index]
    current_api_key_index = (current_api_key_index + 1) % len(API_KEYS)
    return api_key

def call_api(messages, max_tokens=1000, retry_count=3):
    global current_api_key_index
    
    for _ in range(retry_count):
        prompt = ""
        for msg in messages:
            if msg["role"] == "system":
                prompt += f"Instructions: {msg['content']}\n\n"
            else:
                prompt += f"{msg['content']}\n"

        data = {
            "contents": [{
                "parts":[{"text": prompt}]
            }],
            "generationConfig": {
                "maxOutputTokens": max_tokens,
                "temperature": 0.7,
                "topP": 0.95,
            },
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }

        try:
            api_key = get_next_api_key()
            response = requests.post(
                f"{API_URL}?key={api_key}",
                headers={"Content-Type": "application/json"},
                json=data,
                timeout=60
            )
            response_json = response.json()
            
            if response.status_code != 200:
                error_message = response_json.get('error', {}).get('message', 'Unknown error')
                if "quota" in error_message.lower():
                    # Nếu hết quota, thử API key tiếp theo
                    continue
                st.error(f"Lỗi: {error_message}")
                continue

            if 'candidates' in response_json:
                return response_json['candidates'][0]['content']['parts'][0]['text']
                
        except Exception as e:
            st.error(f"Lỗi: {str(e)}")
            continue
    
    return "Không thể tạo nội dung. Vui lòng thử lại sau."

def search_stories(query):
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT story_id, outline, created_at FROM stories
                 WHERE outline LIKE ? OR original_prompt LIKE ?''',
              (f'%{query}%', f'%{query}%'))
    results = c.fetchall()
    conn.close()
    return [{"story_id": r[0], "outline": r[1], "created_at": datetime.fromisoformat(r[2])} for r in results]

def rewrite_story(content, style="normal"):
    styles = {
        "normal": "Viết lại đoạn văn sau với cách diễn đạt mới nhưng giữ nguyên nội dung chính:",
        "creative": "Viết lại đoạn văn sau một cách sáng tạo hơn, thêm các chi tiết mới thú vị:",
        "simple": "Viết lại đoạn văn sau một cách đơn giản, dễ hiểu hơn:",
        "detailed": "Viết lại đoạn văn sau với nhiều chi tiết hơn về cảm xúc và môi trường:"
    }
    
    messages = [
        {"role": "system", "content": "Bạn là một nhà văn chuyên viết truyện ma kinh dị."},
        {"role": "user", "content": f"{styles[style]}\n\n{content}"}
    ]
    return call_api(messages, max_tokens=len(content.split()) + 200)

def generate_story_outline(prompt, num_chapters=10, genre="horror", warnings=None, style=None, custom_genre_guide=None, custom_style_guide=None):
    # Tạo hướng dẫn dựa trên thể loại
    genre_guides = {
        "horror": """Bạn là một nhà văn chuyên viết truyện ma kinh dị. 
        Hãy tạo một khung truyện kinh dị, rùng rợn với những tình tiết giật gân và bầu không khí u ám.""",
        
        "romance": """Bạn là một nhà văn chuyên viết truyện tình cảm lãng mạn.
        Hãy tạo một khung truyện tình cảm sâu sắc, tập trung vào cảm xúc và mối quan hệ giữa các nhân vật.""",
        
        "cultivation": """Bạn là một nhà văn chuyên viết truyện tu tiên, võ hiệp.
        Hãy tạo một khung truyện tu tiên với hệ thống tu luyện rõ ràng, các cấp độ công pháp và thế giới võ lâm huyền ảo.""",
        
        "action": """Bạn là một nhà văn chuyên viết truyện hành động.
        Hãy tạo một khung truyện hành động gay cấn với những pha đánh đấm, rượt đuổi và đối đầu kịch tính.""",
        
        "fantasy": """Bạn là một nhà văn chuyên viết truyện giả tưởng.
        Hãy tạo một khung truyện với thế giới kỳ ảo, phép thuật và sinh vật huyền bí.""",

        "210": """Bạn là một nhà văn chuyên viết truyện 18+.
        Hãy tạo một khung truyện với nội dung 18+, không phù hợp với độc giả dưới 18 tuổi.""",
        
        "custom": custom_genre_guide  # Thêm hướng dẫn tùy chỉnh
    }
    
    # Tạo cảnh báo nội dung
    warning_notes = []
    if warnings:
        if "18+" in warnings:
            warning_notes.append("- Cảnh báo: Truyện có nội dung 18+, không phù hợp với độc giả dưới 18 tuổi")
        if "violence" in warnings:
            warning_notes.append("- Cảnh báo: Truyện có cảnh bạo lực")
        if "horror" in warnings:
            warning_notes.append("- Cảnh báo: Truyện có cảnh kinh dị, rùng rợn")
        if "sensitive" in warnings:
            warning_notes.append("- Cảnh báo: Truyện có nội dung nhạy cảm")
    
    # Tạo hướng dẫn về phong cách
    style_guides = {
        "dark": "Tạo bầu không khí u tối, nặng nề",
        "light": "Tạo bầu không khí nhẹ nhàng, tươi sáng",
        "comedy": "Thêm các yếu tố hài hước",
        "serious": "Giữ giọng văn nghiêm túc, sâu sắc",
        "poetic": "Sử dụng nhiều hình ảnh và ẩn dụ thơ mộng",
        "210": "Tạo bầu không khí lãng mạn, dâm dục và quyến rũ",
        "custom": custom_style_guide  # Thêm phong cách tùy chỉnh
    }
    
    # Lấy hướng dẫn phong cách
    style_note = style_guides.get(style, "")
    if style == "custom" and custom_style_guide:
        style_note = custom_style_guide
    
    # Lấy hướng dẫn thể loại
    genre_guide = genre_guides.get(genre, "")
    if genre == "custom" and custom_genre_guide:
        genre_guide = custom_genre_guide
    
    messages = [
        {"role": "system", "content": f"""{genre_guide}
        
        {style_note}
        
        Hãy tạo một khung truyện chi tiết bằng tiếng Việt, bao gồm:
        1. Tên truyện
        2. Thể loại chính: {genre if genre != "custom" else "Tùy chỉnh"}
        3. Thể loại phụ (nếu có)
        4. Độ tuổi khuyến nghị và cảnh báo nội dung:
        {chr(10).join(warning_notes) if warning_notes else "- Không có cảnh báo đặc biệt"}
        
        5. Giới thiệu ngắn (1-2 đoạn)
        
        6. Nhân vật chính:
           - Tên và vai trò
           - Đặc điểm ngoại hình và tính cách
           - Động lực và mục tiêu
           
        7. Nhân vật phụ:
           - Danh sách nhân vật quan trọng
           - Mối quan hệ với nhân vật chính
           
        8. Bối cảnh:
           - Thời gian và không gian
           - Không khí và màu sắc truyện
           - Quy tắc/Hệ thống thế giới (nếu có)
           
        9. Cốt truyện chính:
           - Điểm khởi đầu
           - Các tình tiết chính
           - Điểm cao trào
           - Kết thúc
           
        10. Số phần dự kiến và nội dung chính của mỗi phần
        
        11. Danh sách {num_chapters} chương cho phần 1:
           (Mỗi chương phải có:
           - Tên chương rõ ràng
            - Tóm tắt nội dung chính 2-3 câu)"""},
        {"role": "user", "content": prompt}
    ]
    return call_api(messages, max_tokens=2000)

def generate_arc_outline(story_outline, arc_number, num_chapters):
    # Tạo outline cho một phần cụ thể của truyện
    messages = [
        {"role": "system", "content": f"""Bạn là một nhà văn chuyên viết truyện ma kinh dị.
        Dựa vào khung truyện sau:
        {story_outline}
        
        Hãy tạo outline chi tiết cho phần {arc_number}, bao gồm:
        1. Tên phần
        2. Mục tiêu của phần này trong cốt truyện tổng thể
        3. Các tình tiết chính cần đạt được
        4. Danh sách {num_chapters} chương:
           (Mỗi chương phải có:
           - Tên chương rõ ràng
           - Tóm tắt nội dung chính 2-3 câu)"""}
    ]
    return call_api(messages, max_tokens=1500)

def generate_chapter(chapter_outline, story_outline, chapter_number, total_chapters, word_count, warnings=None):
    """Tạo nội dung chương với nhiều phong cách và chi tiết hơn"""
    
    # Xác định phong cách và nội dung dựa trên cảnh báo
    content_style = ""
    if warnings:
        if "18+" in warnings:
            content_style += """
            - Thêm các chi tiết về cảnh nóng, quan hệ tình dục một cách tinh tế
            - Miêu tả cảm xúc và ham muốn của nhân vật
            - Sử dụng ngôn từ gợi cảm nhưng không thô tục
            - Tạo không khí lãng mạn và quyến rũ
            """
        if "violence" in warnings:
            content_style += """
            - Thêm các cảnh hành động và bạo lực
            - Miêu tả chi tiết các cuộc đấu tranh
            - Thể hiện sự tàn nhẫn và đau đớn
            """
        if "horror" in warnings:
            content_style += """
            - Tạo không khí kinh dị và rùng rợn
            - Thêm các yếu tố siêu nhiên đáng sợ
            - Miêu tả nỗi sợ hãi và ám ảnh
            """
    
    # Tạo danh sách các kiểu mở đầu đa dạng
    # Tạo danh sách các kiểu mở đầu đa dạng
    opening_styles = [
        "Bắt đầu với một cảnh hành động gay cấn",
        "Mở đầu bằng đối thoại ấn tượng",
        "Khởi đầu với một cảnh tượng bí ẩn",
        "Bắt đầu từ một khoảnh khắc tình cảm",
        "Mở đầu với một cảnh tượng gợi cảm",
        "Khởi đầu từ một giấc mơ hoặc ảo giác",
        "Bắt đầu với một sự kiện bất ngờ",
        "Mở đầu bằng một hồi tưởng",
        "Bắt đầu với một câu hỏi đầy triết lý",
        "Mở đầu theo phong cách báo chí hoặc tài liệu",
        "Bắt đầu bằng một mô tả chi tiết về khung cảnh",
        "Khởi đầu từ góc nhìn của một nhân vật không ngờ tới",
        "Bắt đầu với một tin nhắn hoặc lá thư bí ẩn",
        "Mở đầu bằng một lời tiên tri hoặc cảnh báo",
        "Khởi đầu với một câu nói nội tâm đầy cảm xúc",
        "Bắt đầu với một bức tranh hoặc vật thể đặc biệt",
        "Mở đầu bằng một bản ghi âm hoặc nhật ký",
        "Bắt đầu với một bài thơ hoặc câu hát liên quan đến câu chuyện",
        "Mở đầu bằng một sự kiện lịch sử hoặc giả tưởng",
        "Bắt đầu với một đoạn mô tả thời tiết tạo bầu không khí",
        "Khởi đầu với một giấc mơ hoặc cơn ác mộng",
        "Mở đầu bằng một cuộc trò chuyện điện thoại quan trọng",
        "Bắt đầu với một cảnh sinh hoạt thường ngày của nhân vật chính",
        "Mở đầu với một nhân vật bị truy đuổi",
        "Bắt đầu bằng một nhân vật đang ở trong tình huống nguy cấp",
        "Mở đầu với một câu đố hoặc bí ẩn cần giải quyết",
        "Bắt đầu bằng một lễ hội hoặc sự kiện đông người",
        "Khởi đầu từ một nhân vật đang chạy trốn hoặc giấu giếm điều gì đó",
        "Bắt đầu với một nhân vật tỉnh dậy ở nơi xa lạ",
        "Mở đầu bằng một cú twist ngay từ câu đầu tiên"
    ]

    
    # Chọn ngẫu nhiên kiểu mở đầu
    opening_style = random.choice(opening_styles)
    
    messages = [
        {"role": "system", "content": f"""Bạn là một nhà văn chuyên nghiệp.
        Hãy viết chương {chapter_number}/{total_chapters} với độ dài khoảng {word_count} từ.
        
        Dựa vào khung truyện sau:
        {story_outline}
        
        Phong cách và yêu cầu đặc biệt:
        {content_style}
        
        Yêu cầu chung:
        - {opening_style}
        - Phát triển tình tiết tự nhiên, không gượng ép
        - Xây dựng tâm lý và cảm xúc nhân vật sâu sắc
        - Tạo ra những tình huống bất ngờ nhưng hợp lý
        - Kết hợp hài hòa giữa miêu tả, đối thoại và hành động
        - Sử dụng ngôn ngữ phù hợp với thể loại và đối tượng độc giả
        - Tạo điểm nhấn và cao trào cho chương
        """},
        {"role": "user", "content": f"Viết chương {chapter_number} dựa trên outline: {chapter_outline}"}
    ]
    return call_api(messages, max_tokens=word_count * 2)

def save_story_outline(outline, prompt):
    conn = get_db()
    c = conn.cursor()
    story_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    c.execute('''INSERT INTO stories (story_id, outline, original_prompt, created_at, last_updated)
                 VALUES (?, ?, ?, ?, ?)''', (story_id, outline, prompt, now, now))
    conn.commit()
    conn.close()
    return story_id

def save_arc_outline(story_id, arc_number, outline):
    conn = get_db()
    c = conn.cursor()
    now = datetime.now().isoformat()
    
    # Thêm bảng story_arcs nếu chưa có
    c.execute('''CREATE TABLE IF NOT EXISTS story_arcs
                 (arc_id TEXT PRIMARY KEY, story_id TEXT, arc_number INTEGER,
                  outline TEXT, created_at TIMESTAMP,
                  FOREIGN KEY (story_id) REFERENCES stories(story_id))''')
    
    arc_id = str(uuid.uuid4())
    c.execute('''INSERT INTO story_arcs (arc_id, story_id, arc_number, outline, created_at)
                 VALUES (?, ?, ?, ?, ?)''', (arc_id, story_id, arc_number, outline, now))
    conn.commit()
    conn.close()
    return arc_id

def get_story_list():
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT story_id, outline, created_at FROM stories''')
    results = c.fetchall()
    conn.close()
    return [{"story_id": r[0], "outline": r[1], "created_at": datetime.fromisoformat(r[2])} for r in results]

def save_chapter(story_id, chapter_number, content):
    conn = get_db()
    c = conn.cursor()
    chapter_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    c.execute('''INSERT INTO chapters (chapter_id, story_id, chapter_number, content, created_at)
                 VALUES (?, ?, ?, ?, ?)''', (chapter_id, story_id, chapter_number, content, now))
    c.execute('''UPDATE stories SET last_updated = ? WHERE story_id = ?''', (now, story_id))
    conn.commit()
    conn.close()

def enhance_selected_text(text, enhancement_type):
    prompts = {
        "detail": "Hãy bổ sung thêm chi tiết cho đoạn văn sau (ngắn gọn):",
        "horror": "Hãy tăng cường yếu tố kinh dị cho đoạn văn sau (ngắn gọn):",
        "expand": "Hãy mở rộng đoạn văn sau (ngắn gọn):",
        "dialogue": "Hãy thêm đối thoại vào đoạn văn sau (ngắn gọn):"
    }
    
    messages = [
        {"role": "system", "content": "Bạn là một nhà văn chuyên viết truyện ma kinh dị. Hãy chỉnh sửa ngắn gọn và súc tích."},
        {"role": "user", "content": f"{prompts[enhancement_type]}\n\n{text}"}
    ]
    return call_api(messages, max_tokens=500)

def text_to_speech(text, voice="banmai", speed="", api_key="rNz01K70Q2lG9s2tvF5oGUyQFa16EiwA"):
    """Chuyển đổi text thành speech sử dụng FPT API"""
    url = 'https://api.fpt.ai/hmi/tts/v5'
    headers = {
        'api-key': api_key,
        'speed': speed,
        'voice': voice
    }
    
    try:
        response = requests.post(url, data=text.encode('utf-8'), headers=headers)
        if response.status_code == 200:
            response_data = response.json()
            if 'async' in response_data:
                # Tải file audio về
                audio_url = response_data['async']
                try:
                    # Tạo thư mục audio nếu chưa có
                    os.makedirs('static/audio', exist_ok=True)
                    # Tạo tên file duy nhất
                    audio_filename = f"audio_{uuid.uuid4()}.mp3"
                    audio_path = os.path.join('static/audio', audio_filename)
                    
                    # Thử tải file audio với nhiều lần thử
                    max_retries = 5  # Số lần thử tối đa
                    retry_delay = 5   # Số giây đợi giữa các lần thử
                    
                    for attempt in range(max_retries):
                        try:
                            # Đợi trước khi thử tải
                            time.sleep(retry_delay)
                            
                            # Tải file audio với timeout 30 giây
                            audio_response = requests.get(audio_url, timeout=30)
                            
                            # Kiểm tra kích thước file
                            if audio_response.status_code == 200 and len(audio_response.content) > 0:
                                with open(audio_path, 'wb') as f:
                                    f.write(audio_response.content)
                                # Kiểm tra file đã được tạo thành công
                                if os.path.exists(audio_path) and os.path.getsize(audio_path) > 0:
                                    return True, {'url': audio_url, 'local_path': audio_path}
                            elif audio_response.status_code == 404:
                                if attempt < max_retries - 1:  # Nếu còn lần thử
                                    continue  # Thử lại
                                else:
                                    return False, f"Không thể tải file audio sau {max_retries} lần thử"
                            else:
                                return False, f"Lỗi khi tải file: HTTP {audio_response.status_code}"
                        except requests.Timeout:
                            if attempt < max_retries - 1:  # Nếu còn lần thử
                                continue  # Thử lại
                            else:
                                return False, f"Hết thời gian chờ sau {max_retries} lần thử"
                        except Exception as e:
                            return False, f"Lỗi khi tải audio: {str(e)}"
                    
                    return False, "Không thể tải file audio sau nhiều lần thử"
                except Exception as e:
                    return False, f"Lỗi khi xử lý file audio: {str(e)}"
            return False, "Không nhận được URL âm thanh"
        return False, f"Lỗi API: {response.text}"
    except Exception as e:
        return False, f"Lỗi: {str(e)}"

def get_story_data(story_id):
    conn = get_db()
    c = conn.cursor()
    # Lấy thông tin truyện
    c.execute('''SELECT outline, created_at FROM stories WHERE story_id = ?''', (story_id,))
    story = c.fetchone()
    # Lấy các chương
    c.execute('''SELECT chapter_number, content, created_at FROM chapters
                 WHERE story_id = ? ORDER BY chapter_number''', (story_id,))
    chapters = c.fetchall()
    conn.close()
    
    if story:
        # Lấy tên truyện từ dòng đầu tiên của outline
        title = story[0].split('\n')[0]
        total_chapters = len(chapters)
        is_completed = total_chapters > 0 and total_chapters == int(story[0].split("Danh sách")[1].split("chương")[0].strip())
        
        # Thêm audio_url vào thông tin chương
        chapter_data = []
        for chapter in chapters:
            chapter_info = {
                    "chapter_number": chapter[0],
                    "content": chapter[1],
                "created_at": chapter[2],
                "audio_url": None  # Mặc định không có audio
            }
            chapter_data.append(chapter_info)
        
        return {
            "id": story_id,
            "title": title,
            "outline": story[0],
            "created_at": story[1],
            "total_chapters": total_chapters,
            "is_completed": is_completed,
            "chapters": chapter_data
        }
    return None

def publish_to_web(story_id):
    """Đẩy truyện lên website"""
    story_data = get_story_data(story_id)
    if not story_data:
        return False, "Không tìm thấy truyện"
    
    try:
        # Tạo thư mục static và templates nếu chưa có
        os.makedirs('static', exist_ok=True)
        os.makedirs('templates', exist_ok=True)
        
        # Lưu dữ liệu truyện vào file JSON
        stories_file = 'static/stories.json'
        stories = []
        if os.path.exists(stories_file):
            with open(stories_file, 'r', encoding='utf-8') as f:
                stories = json.load(f)
        
        # Cập nhật hoặc thêm mới truyện
        story_index = next((i for i, s in enumerate(stories) if s['id'] == story_id), -1)
        if story_index >= 0:
            stories[story_index] = story_data
        else:
            stories.append(story_data)
        
        # Lưu lại file JSON
        with open(stories_file, 'w', encoding='utf-8') as f:
            json.dump(stories, f, ensure_ascii=False, indent=2)
        
        return True, "Đã đẩy truyện lên web thành công"
    except Exception as e:
        return False, f"Lỗi khi đẩy truyện lên web: {str(e)}"

def export_to_word(story_id, file_path):
    story_data = get_story_data(story_id)
    if not story_data:
        return False
        
    doc = Document()
    # Thêm tiêu đề
    title = story_data['outline'].split('\n')[0]
    doc.add_heading(title, 0)
    
    # Thêm khung truyện
    doc.add_heading('Khung truyện', level=1)
    doc.add_paragraph(story_data['outline'])
    doc.add_paragraph('\n---\n')
    
    # Thêm các chương
    for chapter in story_data['chapters']:
        doc.add_heading(f'Chương {chapter["chapter_number"]}', level=1)
        doc.add_paragraph(chapter['content'])
        doc.add_paragraph('\n')
    
    doc.save(file_path)
    return True

def export_to_pdf(story_id, file_path):
    story_data = get_story_data(story_id)
    if not story_data:
        return False
        
    pdf = FPDF()
    pdf.add_page()
    
    # Cấu hình font cho tiếng Việt
    try:
        pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
    except:
        # Nếu không có font DejaVu, dùng font mặc định
        pass
    
    pdf.set_font('Arial', '', 12)
    
    # Tiêu đề
    title = story_data['outline'].split('\n')[0]
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, title, ln=True, align='C')
    
    # Khung truyện
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 10, story_data['outline'])
    pdf.ln()
    
    # Các chương
    for chapter in story_data['chapters']:
        pdf.add_page()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, f'Chương {chapter["chapter_number"]}', ln=True)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 10, chapter['content'])
    
    try:
        pdf.output(file_path, 'F')
        return True
    except Exception as e:
        st.error(f"Lỗi khi xuất PDF: {str(e)}")
        return False

def get_chapter_versions(story_id, chapter_number):
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT v.content, v.created_at
                 FROM chapter_versions v
                 JOIN chapters c ON v.chapter_id = c.chapter_id
                 WHERE c.story_id = ? AND c.chapter_number = ?
                 ORDER BY v.created_at DESC''', (story_id, chapter_number))
    versions = c.fetchall()
    conn.close()
    return [{"content": r[0], "created_at": datetime.fromisoformat(r[1])} for r in versions]

def save_chapter_version(story_id, chapter_number, content):
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT chapter_id FROM chapters
                 WHERE story_id = ? AND chapter_number = ?''', (story_id, chapter_number))
    chapter = c.fetchone()
    if chapter:
        version_id = str(uuid.uuid4())
        now = datetime.now().isoformat()
        c.execute('''INSERT INTO chapter_versions (version_id, chapter_id, content, created_at)
                     VALUES (?, ?, ?, ?)''', (version_id, chapter[0], content, now))
        conn.commit()
    conn.close()
    return version_id if chapter else None

def delete_chapter(story_id, chapter_number):
    conn = get_db()
    c = conn.cursor()
    c.execute('''DELETE FROM chapters WHERE story_id = ? AND chapter_number = ?''',
              (story_id, chapter_number))
    conn.commit()
    conn.close()

def delete_story(story_id):
    conn = get_db()
    c = conn.cursor()
    # Xóa tất cả các phiên bản của các chương
    c.execute('''DELETE FROM chapter_versions 
                 WHERE chapter_id IN (
                     SELECT chapter_id FROM chapters WHERE story_id = ?
                 )''', (story_id,))
    # Xóa tất cả các chương
    c.execute('''DELETE FROM chapters WHERE story_id = ?''', (story_id,))
    # Xóa truyện
    c.execute('''DELETE FROM stories WHERE story_id = ?''', (story_id,))
    conn.commit()
    conn.close()

def get_story_chapters(story_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT chapter_number, content, created_at FROM chapters
                 WHERE story_id = ? ORDER BY chapter_number''', (story_id,))
    chapters = c.fetchall()
    conn.close()
    return [{"chapter_number": r[0], "content": r[1], "created_at": datetime.fromisoformat(r[2])} for r in chapters]

def export_chapter_to_word(chapter_data, file_path):
    doc = Document()
    doc.add_heading(f'Chương {chapter_data["chapter_number"]}', 0)
    doc.add_paragraph(chapter_data['content'])
    doc.save(file_path)
    return True

def export_all_chapters_to_word(story_id, file_path):
    story_data = get_story_data(story_id)
    if not story_data:
        return False
        
    doc = Document()
    # Thêm tiêu đề
    title = story_data['outline'].split('\n')[0]
    doc.add_heading(f"Tất cả các chương - {title}", 0)
    
    # Thêm các chương
    for chapter in story_data['chapters']:
        doc.add_heading(f'Chương {chapter["chapter_number"]}', level=1)
        doc.add_paragraph(chapter['content'])
        doc.add_page_break()
    
    doc.save(file_path)
    return True

def auto_generate_chapters(story_id, start_chapter, end_chapter, word_count):
    story_data = get_story_data(story_id)
    if not story_data:
        return False
    
    total_chapters = get_total_chapters_from_outline(story_data['outline'])
    
    # Lấy cảnh báo nội dung từ outline
    warnings = []
    if "18+" in story_data['outline'].lower():
        warnings.append("18+")
    if "bạo lực" in story_data['outline'].lower():
        warnings.append("violence")
    if "kinh dị" in story_data['outline'].lower():
        warnings.append("horror")
    
    for chapter_number in range(start_chapter, end_chapter + 1):
        if chapter_number <= total_chapters:
            # Tạo nội dung chương với cảnh báo
            chapter_content = generate_chapter(
                f"Viết chương {chapter_number}",
                story_data['outline'],
                chapter_number,
                total_chapters,
                word_count,
                warnings=warnings
            )
            # Lưu chương
            save_chapter(story_id, chapter_number, chapter_content)
    return True

def get_total_chapters_from_outline(outline):
    """Lấy tổng số chương từ outline với xử lý lỗi"""
    try:
        # Tìm phần "Danh sách X chương"
        if not outline or "Danh sách" not in outline:
            return 10  # Giá trị mặc định nếu không tìm thấy
        
        # Tìm số chương bằng regex
        import re
        
        # Thử tìm theo mẫu "Danh sách X chương"
        matches = re.findall(r'Danh sách\s+(\d+)\s+chương', outline)
        if matches:
            return int(matches[0])
            
        # Thử tìm theo mẫu "X chương"
        matches = re.findall(r'(\d+)\s+chương', outline)
        if matches:
            return int(matches[0])
            
        # Thử tìm theo mẫu "Chương X:"
        chapter_numbers = re.findall(r'Chương\s+(\d+):', outline)
        if chapter_numbers:
            return max(map(int, chapter_numbers))
        
        # Đếm số lần xuất hiện của từ "Chương"
        chapter_count = len(re.findall(r'Chương\s+\d+', outline, re.IGNORECASE))
        if chapter_count > 0:
            return chapter_count
        
        return 10  # Giá trị mặc định nếu không tìm được số
    except Exception as e:
        st.error(f"Lỗi khi đọc số chương: {str(e)}")
        return 10  # Giá trị mặc định nếu có lỗi

def get_story_arcs(story_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT arc_number, outline, created_at FROM story_arcs
                 WHERE story_id = ? ORDER BY arc_number''', (story_id,))
    arcs = c.fetchall()
    conn.close()
    return [{"arc_number": r[0], "outline": r[1], "created_at": datetime.fromisoformat(r[2])} for r in arcs]

def generate_long_chapter(chapter_outline, story_outline, chapter_number, total_chapters, min_words=1000):
    """Tạo nội dung chương dài bằng cách chia thành nhiều phần và ghép lại"""
    
    # Chia thành 3 phần: mở đầu, thân truyện, kết thúc
    parts = []
    
    # Phần 1: Mở đầu (khoảng 20% nội dung)
    intro_messages = [
        {"role": "system", "content": f"""Bạn là một nhà văn chuyên viết truyện ma kinh dị.
        Hãy viết phần mở đầu của chương {chapter_number}/{total_chapters}.
        Dựa vào khung truyện sau:
        {story_outline}
        
        Yêu cầu:
        - Giới thiệu bối cảnh và tình huống
        - Tạo không khí rùng rợn
        - Độ dài khoảng {min_words//5} từ"""}
    ]
    intro = call_api(intro_messages, max_tokens=1000)
    parts.append(intro)
    
    # Phần 2: Thân truyện (khoảng 60% nội dung)
    main_messages = [
        {"role": "system", "content": f"""Tiếp tục viết phần thân truyện của chương {chapter_number}, sau phần mở đầu:
        {intro}
        
        Dựa vào outline:
        {chapter_outline}
        
        Yêu cầu:
        - Phát triển tình tiết chính
        - Tạo căng thẳng và kịch tính
        - Miêu tả chi tiết cảm xúc nhân vật
        - Độ dài khoảng {min_words*3//5} từ"""}
    ]
    main_content = call_api(main_messages, max_tokens=2000)
    parts.append(main_content)
    
    # Phần 3: Kết thúc (khoảng 20% nội dung)
    ending_messages = [
        {"role": "system", "content": f"""Viết phần kết thúc của chương {chapter_number}, sau phần:
        {main_content}
        
        Yêu cầu:
        - Tạo điểm nhấn hoặc twist
        - Kết nối với chương tiếp theo
        - Duy trì không khí rùng rợn
        - Độ dài khoảng {min_words//5} từ"""}
    ]
    ending = call_api(ending_messages, max_tokens=1000)
    parts.append(ending)
    
    # Ghép các phần lại
    full_chapter = "\n\n".join(parts)
    
    # Kiểm tra độ dài và tạo thêm nội dung nếu cần
    while len(full_chapter.split()) < min_words:
        expand_messages = [
            {"role": "system", "content": f"""Mở rộng đoạn văn sau, thêm chi tiết và miêu tả:
            {full_chapter}
            
            Yêu cầu:
            - Giữ nguyên cốt truyện
            - Thêm chi tiết miêu tả và đối thoại
            - Tăng độ dài thêm {min_words - len(full_chapter.split())} từ"""}
        ]
        expanded = call_api(expand_messages, max_tokens=1000)
        full_chapter = expanded
    
    return full_chapter

def main():
    global API_KEYS
    st.set_page_config(layout="wide")
    st.title("Công Cụ Tạo Truyện Tự Động 📚")
    
    menu = st.sidebar.selectbox(
        "Chọn chức năng",
        ["Tạo Truyện Mới", "Tìm Kiếm Truyện", "Viết & Chỉnh Sửa", "Quản Lý Truyện", "Quản Lý Audio"]
    )
    
    if menu == "Quản Lý Audio":
        st.header("Quản Lý Audio")
        
        # Phần cấu hình API FPT
        st.subheader("Cấu hình API FPT")
        api_keys_str = st.text_area(
            "Nhập các API key FPT (mỗi key một dòng):",
            value="rNz01K70Q2lG9s2tvF5oGUyQFa16EiwA",  # API key mặc định
            help="Thêm nhiều API key để tăng khả năng chuyển đổi"
        )
        fpt_api_keys = [key.strip() for key in api_keys_str.split('\n') if key.strip()]
        
        # Chọn truyện để tạo audio
        stories = get_story_list()
        if not stories:
            st.info("Chưa có truyện nào được tạo")
        else:
            selected_story = st.selectbox(
                "Chọn truyện để tạo audio:",
                options=[story['story_id'] for story in stories],
                format_func=lambda x: next(s['outline'].split('\n')[0] for s in stories if s['story_id'] == x)
            )
            
            if selected_story:
                story_data = get_story_data(selected_story)
                chapters = get_story_chapters(selected_story)
                
                # Cấu hình giọng đọc
                col1, col2 = st.columns(2)
                with col1:
                    voice = st.selectbox(
                        "Chọn giọng đọc:",
                        ["banmai", "thuminh", "leminh", "myan", "lannhi", "linhsan"]
                    )
                with col2:
                    speed = st.slider(
                        "Tốc độ đọc:",
                        min_value=-3,
                        max_value=3,
                        value=0
                    )
                
                # Hiển thị danh sách chương và nút tạo audio
                st.subheader(f"Danh sách chương - {story_data['title']}")
                
                for chapter in chapters:
                    with st.expander(f"Chương {chapter['chapter_number']}"):
                        st.write(chapter['content'][:200] + "...")  # Hiển thị preview nội dung
                        
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            # Kiểm tra xem đã có audio chưa
                            if 'audio_url' in chapter and chapter['audio_url']:
                                st.markdown(f"""
                                <audio controls>
                                    <source src="{chapter['audio_url']}" type="audio/mp3">
                                    Trình duyệt của bạn không hỗ trợ audio.
                                </audio>
                                """, unsafe_allow_html=True)
                            else:
                                st.info("Chưa có audio")
                        
                        with col2:
                            if st.button("Tạo Audio", key=f"create_audio_{chapter['chapter_number']}"):
                                with st.spinner('Đang tạo audio...'):
                                    # Thử với từng API key cho đến khi thành công
                                    for api_key in fpt_api_keys:
                                        success, result = text_to_speech(
                                            chapter['content'],
                                            voice=voice,
                                            speed=str(speed),
                                            api_key=api_key
                                        )
                                        if success:
                                            st.success("Đã tạo audio thành công!")
                                            # Hiển thị audio player
                                            st.markdown(f"""
                                            <audio controls>
                                                <source src="{result['url']}" type="audio/mp3">
                                                Trình duyệt của bạn không hỗ trợ audio.
                                            </audio>
                                            """, unsafe_allow_html=True)
                                            
                                            # Thêm nút tải về
                                            with open(result['local_path'], 'rb') as f:
                                                st.download_button(
                                                    label="Tải audio về máy",
                                                    data=f,
                                                    file_name=f"chuong_{chapter['chapter_number']}_audio.mp3",
                                                    mime="audio/mp3"
                                                )
                                            
                                            # Lưu URL audio vào database
                                            save_audio_url(selected_story, chapter['chapter_number'], result)
                                            break
                                        else:
                                            st.error(f"Lỗi với API key {api_key[:10]}...: {result}")
                                            continue
    
    elif menu == "Tạo Truyện Mới":
        st.header("Tạo Truyện Mới")
        
        # Tạo layout 2 cột cho các tùy chọn
        col1, col2 = st.columns(2)
        
        with col1:
            # Chọn thể loại chính
            genre_options = {
                "horror": "Truyện Ma - Kinh Dị",
                "romance": "Tình Cảm - Lãng Mạn",
                "cultivation": "Tu Tiên - Võ Hiệp",
                "action": "Hành Động - Phiêu Lưu",
                "fantasy": "Giả Tưởng - Kỳ Ảo",
                "210": "18+ - 210 :)))",
                "custom": "Thể Loại Tùy Chỉnh"
            }
            genre = st.selectbox(
                "Chọn thể loại:",
                options=list(genre_options.keys()),
                format_func=lambda x: genre_options[x]
            )
            
            # Nếu chọn thể loại tùy chỉnh
            custom_genre_guide = None
            if genre == "custom":
                custom_genre_name = st.text_input("Tên thể loại mới:")
                custom_genre_guide = st.text_area(
                    "Mô tả hướng dẫn cho thể loại:",
                    help="Mô tả chi tiết về đặc điểm, yêu cầu và phong cách của thể loại này"
                )
            
            # Chọn phong cách viết
            style_options = {
                "dark": "U tối, nặng nề",
                "light": "Nhẹ nhàng, tươi sáng",
                "comedy": "Hài hước",
                "serious": "Nghiêm túc",
                "poetic": "Thơ mộng",
                "210": "18+ - 210 :)))",
                "custom": "Phong Cách Tùy Chỉnh"
            }
            style = st.selectbox(
                "Chọn phong cách:",
                options=list(style_options.keys()),
                format_func=lambda x: style_options[x]
            )
            
            # Nếu chọn phong cách tùy chỉnh
            custom_style_guide = None
            if style == "custom":
                custom_style_name = st.text_input("Tên phong cách mới:")
                custom_style_guide = st.text_area(
                    label="Mô tả hướng dẫn cho phong cách:",
                    help="Mô tả chi tiết về cách viết, giọng văn và không khí của phong cách này"
                )
        
        with col2:
            # Chọn các cảnh báo nội dung
            warnings = st.multiselect(
                "Cảnh báo nội dung:",
                ["18+", "violence", "horror", "sensitive"],
                format_func=lambda x: {
                    "18+": "Nội dung 18+",
                    "violence": "Bạo lực",
                    "horror": "Kinh dị",
                    "sensitive": "Nội dung nhạy cảm"
                }[x]
            )
            
            # Cấu hình độ dài truyện
            num_chapters = st.number_input("Số chương:", min_value=1, max_value=2000, value=10)
            words_per_chapter = st.number_input("Số từ mỗi chương:", min_value=1000, max_value=10000, value=2000, step=500)
        
        # Thêm phần cấu hình API keys
        with st.expander("Cấu hình API"):
            api_keys_str = st.text_area(
                "Nhập các API key (mỗi key một dòng):",
                value="\n".join(API_KEYS),
                help="Thêm nhiều API key để tăng khả năng tạo nội dung dài"
            )
            if api_keys_str:
                API_KEYS[:] = [key.strip() for key in api_keys_str.split('\n') if key.strip()]
                st.success(f"Đã cập nhật {len(API_KEYS)} API key")

        # Nhập ý tưởng truyện
        prompt = st.text_area(
            "Nhập ý tưởng cho truyện của bạn:",
            height=150,
            help="Mô tả ý tưởng, bối cảnh, nhân vật, không khí truyện bạn muốn tạo"
        )
        
        if st.button("Tạo Khung Truyện"):
            if prompt:
                with st.spinner('Đang tạo khung truyện...'):
                    outline = generate_story_outline(
                        prompt,
                        num_chapters=num_chapters,
                        genre=genre,
                        warnings=warnings,
                        style=style,
                        custom_genre_guide=custom_genre_guide,
                        custom_style_guide=custom_style_guide
                    )
                    story_id = save_story_outline(outline, prompt)
                    st.markdown("### Khung Truyện:")
                    st.write(outline)
                    st.success(f"Đã lưu truyện với ID: {story_id}")
            else:
                st.warning("Vui lòng nhập ý tưởng cho truyện!")

    elif menu == "Tìm Kiếm Truyện":
        st.header("Tìm Kiếm Truyện")
        search_query = st.text_input("Nhập từ khóa tìm kiếm:")
        
        if search_query:
            stories = search_stories(search_query)
            if stories:
                st.success(f"Tìm thấy {len(stories)} kết quả")
                for story in stories:
                    with st.expander(f"Truyện tạo ngày {story['created_at'].strftime('%d/%m/%Y %H:%M')}"):
                        st.write(story['outline'])
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("Chọn để chỉnh sửa", key=f"edit_{story['story_id']}"):
                                st.session_state['current_story'] = story['story_id']
                                st.session_state['current_outline'] = story['outline']
                                st.rerun()
                        with col2:
                            # Xử lý viết lại truyện
                            rewrite_button = st.button("Viết lại", key=f"rewrite_{story['story_id']}")
                            if rewrite_button:
                                style = st.selectbox(
                                    "Chọn phong cách viết lại:",
                                    ["normal", "creative", "simple", "detailed"],
                                    format_func=lambda x: {
                                        "normal": "Bình thường",
                                        "creative": "Sáng tạo", 
                                        "simple": "Đơn giản",
                                        "detailed": "Chi tiết"
                                    }[x],
                                    key=f"style_{story['story_id']}"
                                )
                                
                                confirm_button = st.button("Xác nhận viết lại", key=f"confirm_rewrite_{story['story_id']}")
                                if confirm_button:
                                    with st.spinner('Đang viết lại truyện...'):
                                        rewritten = rewrite_story(story['outline'], style)
                                        st.markdown("### Phiên bản viết lại:")
                                        st.write(rewritten)
            else:
                st.info("Không tìm thấy truyện nào phù hợp")

    elif menu == "Viết & Chỉnh Sửa":
        st.header("Viết & Chỉnh Sửa Truyện")
        
        if 'current_story' in st.session_state and 'current_outline' in st.session_state:
            st.subheader("Khung Truyện")
            current_outline = st.session_state['current_outline']
            st.write(current_outline)
            
            # Lấy tổng số chương
            total_chapters = get_total_chapters_from_outline(current_outline)
            st.info(f"Tổng số chương: {total_chapters}")
            
            # Thêm phần quản lý arc
            st.subheader("Quản lý phần truyện")
            arcs = get_story_arcs(st.session_state['current_story'])
            
            col1, col2 = st.columns(2)
            with col1:
                next_arc = len(arcs) + 1
                num_chapters = st.number_input("Số chương cho phần mới:", min_value=1, max_value=50, value=10)
                if st.button("Tạo phần mới"):
                    with st.spinner(f'Đang tạo outline cho phần {next_arc}...'):
                        arc_outline = generate_arc_outline(st.session_state['current_outline'], next_arc, num_chapters)
                        arc_id = save_arc_outline(st.session_state['current_story'], next_arc, arc_outline)
                        st.success(f"Đã tạo phần {next_arc}")
                        st.rerun()
            
            with col2:
                if arcs:
                    st.write(f"Đã có {len(arcs)} phần")
                    for arc in arcs:
                        with st.expander(f"Phần {arc['arc_number']}"):
                            st.write(arc['outline'])
            
            st.markdown("---")
            
            # Thêm nút xuất file
            col_export1, col_export2, col_export3 = st.columns(3)
            with col_export1:
                if st.button("Xuất toàn bộ truyện (Word)"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        if export_to_word(st.session_state['current_story'], tmp.name):
                            with open(tmp.name, 'rb') as f:
                                st.download_button(
                                    "Tải xuống truyện hoàn chỉnh",
                                    f,
                                    file_name="truyen_ma_full.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
            
            with col_export2:
                if st.button("Xuất tất cả các chương (Word)"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        if export_all_chapters_to_word(st.session_state['current_story'], tmp.name):
                            with open(tmp.name, 'rb') as f:
                                st.download_button(
                                    "Tải xuống tất cả chương",
                                    f,
                                    file_name="tat_ca_chuong.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
            
            with col_export3:
                if st.button("Xuất file PDF"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                        if export_to_pdf(st.session_state['current_story'], tmp.name):
                            with open(tmp.name, 'rb') as f:
                                st.download_button(
                                    "Tải xuống file PDF",
                                    f,
                                    file_name="truyen_ma.pdf",
                                    mime="application/pdf"
                                )
            
            # Thêm chức năng tự động tạo nhiều chương
            st.subheader("Tự động tạo chương")
            col_auto1, col_auto2, col_auto3 = st.columns(3)
            with col_auto1:
                start_chapter = st.number_input("Từ chương:", min_value=1, value=1)
            with col_auto2:
                end_chapter = st.number_input("Đến chương:", min_value=1, value=5)
            with col_auto3:
                auto_word_count = st.number_input("Số từ mỗi chương:", min_value=500, max_value=5000, value=1000, step=100)
            
            if st.button("Tự động tạo các chương"):
                total_chapters = get_total_chapters_from_outline(current_outline)
                if end_chapter > total_chapters:
                    st.error(f"Số chương tối đa là {total_chapters}")
                elif start_chapter > end_chapter:
                    st.error("Chương bắt đầu phải nhỏ hơn chương kết thúc")
                else:
                    with st.spinner(f'Đang tạo các chương từ {start_chapter} đến {end_chapter}...'):
                        if auto_generate_chapters(
                            st.session_state['current_story'],
                            start_chapter,
                            end_chapter,
                            auto_word_count
                        ):
                            st.success("Đã tạo xong các chương!")
                            st.rerun()
                        else:
                            st.error("Có lỗi xảy ra khi tạo chương")
            
            # Hiển thị danh sách các chương đã viết
            st.subheader("Các chương đã viết")
            chapters = get_story_chapters(st.session_state['current_story'])
            for chapter in chapters:
                with st.expander(f"Chương {chapter['chapter_number']} - {chapter['created_at'].strftime('%d/%m/%Y %H:%M')}"):
                    st.write(chapter['content'])
                    col1, col2, col3, col4, col5 = st.columns(5)
                    with col1:
                        if st.button("Xóa", key=f"del_{chapter['chapter_number']}"):
                            delete_chapter(st.session_state['current_story'], chapter['chapter_number'])
                            st.success("Đã xóa chương!")
                            st.rerun()
                    with col2:
                        if st.button("Viết lại", key=f"rewrite_{chapter['chapter_number']}"):
                            new_content = rewrite_story(chapter['content'])
                            save_chapter_version(st.session_state['current_story'], chapter['chapter_number'], new_content)
                            st.success("Đã tạo phiên bản mới!")
                    with col3:
                        if st.button("Xem phiên bản", key=f"versions_{chapter['chapter_number']}"):
                            versions = get_chapter_versions(st.session_state['current_story'], chapter['chapter_number'])
                            for version in versions:
                                st.text(f"Phiên bản {version['created_at'].strftime('%d/%m/%Y %H:%M')}")
                                st.text_area("Nội dung:", version['content'], height=200, key=f"v_{version['version_id']}")
                    with col4:
                        if st.button("Chỉnh sửa", key=f"edit_{chapter['chapter_number']}"):
                            st.session_state['editing_chapter'] = chapter['chapter_number']
                            st.session_state['editing_content'] = chapter['content']
                    with col5:
                        if st.button("Xuất Word", key=f"export_{chapter['chapter_number']}"):
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                if export_chapter_to_word(chapter, tmp.name):
                                    with open(tmp.name, 'rb') as f:
                                        st.download_button(
                                            f"Tải Chương {chapter['chapter_number']}",
                                            f,
                                            file_name=f"chuong_{chapter['chapter_number']}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_{chapter['chapter_number']}"
                                        )
            
            # Hiển thị tổng quan về tiến độ
            st.markdown("---")
            st.subheader("Tiến độ truyện")
            progress = len(chapters) / total_chapters
            st.progress(progress)
            st.write(f"Đã viết: {len(chapters)}/{total_chapters} chương")
            
            st.markdown("---")
            
            # Khu vực viết và chỉnh sửa
            col1, col2 = st.columns([6, 4])
            
            with col1:
                st.subheader("Viết Chương Mới")
                next_chapter = len(chapters) + 1
                if next_chapter <= total_chapters:
                    chapter_number = st.number_input("Số chương:", min_value=1, max_value=total_chapters, value=next_chapter)
                    word_count = st.number_input("Số từ:", min_value=500, max_value=5000, value=1000, step=100)
                    
                    if st.button("Tạo nội dung chương"):
                        chapter_content = generate_chapter(
                            f"Viết chương {chapter_number}",
                            st.session_state['current_outline'],
                            chapter_number,
                            total_chapters,
                            word_count
                        )
                        st.session_state['current_chapter'] = chapter_content
                        st.write(chapter_content)
                    
                    chapter_content = st.text_area(
                        "Nội dung chương:",
                        value=st.session_state.get('current_chapter', ''),
                        height=400
                    )
                    
                    if st.button("Lưu chương"):
                        if 'editing_chapter' in st.session_state:
                            # Lưu phiên bản mới khi đang chỉnh sửa
                            save_chapter_version(st.session_state['current_story'], st.session_state['editing_chapter'], chapter_content)
                            del st.session_state['editing_chapter']
                        else:
                            # Lưu chương mới
                            save_chapter(st.session_state['current_story'], chapter_number, chapter_content)
                        st.success("Đã lưu chương thành công!")
                        st.rerun()
                else:
                    st.warning("Đã hoàn thành tất cả các chương!")

        else:
            st.info("Vui lòng chọn một truyện từ Tìm Kiếm Truyện để bắt đầu chỉnh sửa")

    elif menu == "Quản Lý Truyện":
        st.header("Quản Lý Truyện")
        stories = get_story_list()
        
        if not stories:
            st.info("Chưa có truyện nào được tạo")
        else:
            st.success(f"Có {len(stories)} truyện")
            
            # Tạo bảng hiển thị
            col_titles = st.columns([3, 2, 1, 1, 1, 1, 1, 1])
            col_titles[0].markdown("### Tên truyện")
            col_titles[1].markdown("### Ngày tạo")
            col_titles[2].markdown("### Số chương")
            col_titles[3].markdown("### Xóa")
            col_titles[4].markdown("### Xuất Word")
            col_titles[5].markdown("### Xuất chương")
            col_titles[6].markdown("### Đẩy web")
            col_titles[7].markdown("### Audio")
            
            st.markdown("---")
            
            for story in stories:
                # Khởi tạo session state cho xóa truyện
                delete_key = f"delete_{story['story_id']}"
                if delete_key not in st.session_state:
                    st.session_state[delete_key] = False
                
                cols = st.columns([3, 2, 1, 1, 1, 1, 1, 1])
                
                # Lấy thông tin chi tiết truyện
                story_data = get_story_data(story['story_id'])
                title = story_data['title']
                total_chapters = story_data['total_chapters']
                
                # Cột tên truyện
                with cols[0]:
                    st.markdown(f"**{title}**")
                    if st.button("Xem chi tiết", key=f"view_{story['story_id']}"):
                        st.markdown("#### Chi tiết truyện:")
                        st.write(story['outline'])
                
                # Cột ngày tạo
                with cols[1]:
                    st.write(story['created_at'].strftime('%d/%m/%Y %H:%M'))
                
                # Cột số chương
                with cols[2]:
                    st.write(f"{total_chapters} chương")
                
                # Cột xóa truyện
                with cols[3]:
                    # Hiển thị nút xóa và xác nhận
                    if not st.session_state[delete_key]:
                        if st.button("🗑️", key=f"del_btn_{story['story_id']}"):
                            st.session_state[delete_key] = True
                    else:
                        st.warning("Xác nhận xóa?")
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("✔️ Có", key=f"confirm_{story['story_id']}"):
                                delete_story(story['story_id'])
                                st.success("Đã xóa truyện!")
                                st.rerun()
                        with col2:
                            if st.button("❌ Không", key=f"cancel_{story['story_id']}"):
                                st.session_state[delete_key] = False
                                st.rerun()
                
                # Cột xuất toàn bộ Word
                with cols[4]:
                    if st.button("📄", key=f"export_full_{story['story_id']}"):
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            if export_to_word(story['story_id'], tmp.name):
                                with open(tmp.name, 'rb') as f:
                                    st.download_button(
                                        "Tải xuống",
                                        f,
                                        file_name=f"{title}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"download_full_{story['story_id']}"
                                    )
                
                # Cột xuất từng chương
                with cols[5]:
                    chapters = get_story_chapters(story['story_id'])
                    if chapters:
                        if st.button("📑", key=f"export_chapters_{story['story_id']}"):
                            st.markdown("#### Tải xuống từng chương:")
                            for chapter in chapters:
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                    if export_chapter_to_word(chapter, tmp.name):
                                        with open(tmp.name, 'rb') as f:
                                            st.download_button(
                                                f"Chương {chapter['chapter_number']}",
                                                f,
                                                file_name=f"{title}_chuong_{chapter['chapter_number']}.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                key=f"download_chapter_{story['story_id']}_{chapter['chapter_number']}"
                                            )
                
                # Cột đẩy lên web
                with cols[6]:
                    if st.button("🌐", key=f"publish_{story['story_id']}"):
                        success, message = publish_to_web(story['story_id'])
                        if success:
                            st.success(message)
                        else:
                            st.error(message)
                
                # Cột tạo audio
                with cols[7]:
                    if st.button("🔊", key=f"audio_{story['story_id']}"):
                        chapters = get_story_chapters(story['story_id'])
                        if chapters:
                            st.markdown("#### Tạo audio cho từng chương:")
                            # Chọn giọng đọc
                            voice = st.selectbox(
                                "Chọn giọng đọc:",
                                ["banmai", "thuminh", "leminh", "myan", "lannhi", "linhsan"],
                                key=f"voice_{story['story_id']}"
                            )
                            # Chọn tốc độ đọc
                            speed = st.slider(
                                "Tốc độ đọc:",
                                min_value=-3,
                                max_value=3,
                                value=0,
                                key=f"speed_{story['story_id']}"
                            )
                            
                            for chapter in chapters:
                                success, result = text_to_speech(
                                    chapter['content'],
                                    voice=voice,
                                    speed=str(speed)
                                )
                                if success:
                                    st.markdown(f"""
                                    ##### Chương {chapter['chapter_number']}
                                    <audio controls>
                                        <source src="{result['url']}" type="audio/mp3">
                                        Trình duyệt của bạn không hỗ trợ audio.
                                    </audio>
                                    """, unsafe_allow_html=True)
                                else:
                                    st.error(f"Lỗi tạo audio cho chương {chapter['chapter_number']}: {result}")
                
                st.markdown("---")

    elif menu == "Quản Lý Audio":
        st.header("Quản Lý Audio")
        
        # Phần cấu hình API FPT
        st.subheader("Cấu hình API FPT")
        api_keys_str = st.text_area(
            "Nhập các API key FPT (mỗi key một dòng):",
            value="rNz01K70Q2lG9s2tvF5oGUyQFa16EiwA",  # API key mặc định
            help="Thêm nhiều API key để tăng khả năng chuyển đổi"
        )
        fpt_api_keys = [key.strip() for key in api_keys_str.split('\n') if key.strip()]
        
        # Chọn truyện để tạo audio
        stories = get_story_list()
        if not stories:
            st.info("Chưa có truyện nào được tạo")
        else:
            selected_story = st.selectbox(
                "Chọn truyện để tạo audio:",
                options=[story['story_id'] for story in stories],
                format_func=lambda x: next(s['outline'].split('\n')[0] for s in stories if s['story_id'] == x)
            )
            
            if selected_story:
                story_data = get_story_data(selected_story)
                chapters = get_story_chapters(selected_story)
                
                # Cấu hình giọng đọc
                col1, col2 = st.columns(2)
                with col1:
                    voice = st.selectbox(
                        "Chọn giọng đọc:",
                        ["banmai", "thuminh", "leminh", "myan", "lannhi", "linhsan"]
                    )
                with col2:
                    speed = st.slider(
                        "Tốc độ đọc:",
                        min_value=-3,
                        max_value=3,
                        value=0
                    )
                
                # Hiển thị danh sách chương và nút tạo audio
                st.subheader(f"Danh sách chương - {story_data['title']}")
                
                for chapter in chapters:
                    with st.expander(f"Chương {chapter['chapter_number']}"):
                        st.write(chapter['content'][:200] + "...")  # Hiển thị preview nội dung
                        
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            # Kiểm tra xem đã có audio chưa
                            if 'audio_url' in chapter and chapter['audio_url']:
                                st.markdown(f"""
                                <audio controls>
                                    <source src="{chapter['audio_url']}" type="audio/mp3">
                                    Trình duyệt của bạn không hỗ trợ audio.
                                </audio>
                                """, unsafe_allow_html=True)
                            else:
                                st.info("Chưa có audio")
                        
                        with col2:
                            if st.button("Tạo Audio", key=f"create_audio_{chapter['chapter_number']}"):
                                with st.spinner('Đang tạo audio...'):
                                    # Thử với từng API key cho đến khi thành công
                                    for api_key in fpt_api_keys:
                                        success, result = text_to_speech(
                                            chapter['content'],
                                            voice=voice,
                                            speed=str(speed),
                                            api_key=api_key
                                        )
                                        if success:
                                            st.success("Đã tạo audio thành công!")
                                            # Hiển thị audio player
                                            st.markdown(f"""
                                            <audio controls>
                                                <source src="{result['url']}" type="audio/mp3">
                                                Trình duyệt của bạn không hỗ trợ audio.
                                            </audio>
                                            """, unsafe_allow_html=True)
                                            
                                            # Thêm nút tải về
                                            with open(result['local_path'], 'rb') as f:
                                                st.download_button(
                                                    label="Tải audio về máy",
                                                    data=f,
                                                    file_name=f"chuong_{chapter['chapter_number']}_audio.mp3",
                                                    mime="audio/mp3"
                                                )
                                            
                                            # Lưu URL audio vào database
                                            save_audio_url(selected_story, chapter['chapter_number'], result)
                                            break
                                        else:
                                            st.error(f"Lỗi với API key {api_key[:10]}...: {result}")
                                            continue

def save_audio_url(story_id, chapter_number, audio_data):
    """Lưu URL audio và đường dẫn local vào database"""
    conn = get_db()
    c = conn.cursor()
    
    # Thêm các cột audio nếu chưa có
    try:
        c.execute('ALTER TABLE chapters ADD COLUMN audio_url TEXT')
        c.execute('ALTER TABLE chapters ADD COLUMN audio_local_path TEXT')
    except sqlite3.OperationalError:
        pass  # Cột đã tồn tại
    
    # Cập nhật URL audio và đường dẫn local
    c.execute('''UPDATE chapters 
                 SET audio_url = ?, audio_local_path = ?
                 WHERE story_id = ? AND chapter_number = ?''',
              (audio_data['url'], audio_data['local_path'], story_id, chapter_number))
    conn.commit()
    conn.close()

if __name__ == "__main__":
    main() 